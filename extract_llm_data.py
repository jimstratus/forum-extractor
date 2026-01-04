import os
import re
import shutil
from pathlib import Path
import json
from datetime import datetime
import traceback

# Import libraries for different file formats
try:
    import docx  # For .docx files
    from PyPDF2 import PdfReader  # For .pdf files
    import pandas as pd  # For .xlsx, .csv files
    from bs4 import BeautifulSoup  # For HTML/XML content if needed
except ImportError as e:
    print(f"Error importing required libraries: {e}")
    print("Please install required libraries using: pip install python-docx PyPDF2 beautifulsoup4 pandas openpyxl")
    exit(1)

# Define paths relative to the current working directory
BASE_DIR = Path(os.getcwd())
SOURCE_NOVELS_DIR = BASE_DIR / "EOTIR Novels"
SOURCE_RPG_DIR = BASE_DIR / "EOTIR RPG"
TARGET_DIR = BASE_DIR / "LLM"

# Define file type categories
CHARACTER_PROFILE_PATTERNS = [
    r'char.*profile', r'personnel.*file', r'biography', 
    r'charac.*of', r'dramatis.*personae'
]

IN_GAME_DOCS_PATTERNS = {
    'laws_policies': [r'act', r'law', r'policy', r'charter', r'governance', r'declaration'],
    'intelligence': [r'intelligence', r'report', r'brief', r'security', r'field.*op'],
    'technical': [r'tech', r'specification', r'ship.*name', r'oversector', r'map'],
    'diplomatic': [r'treaty', r'diplomatic', r'announcement', r'peace', r'order', r'declaration', r'sovereign']
}

NARRATIVE_PATTERNS = [r'chapter', r'book', r'scenario', r'sequence', r'event', r'beginning']
WORLDBUILDING_PATTERNS = [r'world', r'timeline', r'history', r'lore', r'ir\s*ship', r'galaxy', r'map']
OOC_PATTERNS = {
    'rules': [r'rule', r'guideline'],
    'admin': [r'admin', r'log', r'announcement', r'change'],
    'templates': [r'template', r'form', r'structure']
}
TIMELINE_PATTERNS = [r'timeline', r'sequence.*event', r'history']

# Dictionary to track processed files and their categories
processed_files = {}
file_counts = {
    "character_profiles": 0,
    "in_game_documents": {
        "laws_policies": 0,
        "intelligence": 0,
        "technical": 0,
        "diplomatic": 0
    },
    "narratives": {
        "completed_scenarios": 0,
        "unfinished_scenarios": 0
    },
    "worldbuilding": 0,
    "ooc_content": {
        "rules": 0,
        "admin": 0,
        "templates": 0
    },
    "timeline": 0,
    "unknown": 0,
    "skipped": 0,
    "error": 0
}

# Functions to extract text from different file formats
def extract_text_from_docx(file_path):
    """Extract text from a .docx file."""
    try:
        doc = docx.Document(file_path)
        full_text = []
        
        # Add document properties as metadata if available
        properties = {}
        try:
            core_properties = doc.core_properties
            if hasattr(core_properties, 'title') and core_properties.title:
                properties['title'] = core_properties.title
            if hasattr(core_properties, 'author') and core_properties.author:
                properties['author'] = core_properties.author
            if hasattr(core_properties, 'created') and core_properties.created:
                properties['created'] = str(core_properties.created)
            if hasattr(core_properties, 'modified') and core_properties.modified:
                properties['modified'] = str(core_properties.modified)
        except:
            pass
        
        # Extract paragraphs
        for para in doc.paragraphs:
            if para.text.strip():
                full_text.append(para.text)
        
        # Extract tables if present
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_text.append(cell.text.strip())
                if row_text:
                    full_text.append(" | ".join(row_text))
        
        return {"text": "\n".join(full_text), "metadata": properties}
    except Exception as e:
        print(f"Error extracting text from {file_path}: {e}")
        return {"text": f"ERROR: Could not extract text from {file_path}", "metadata": {}}

def extract_text_from_pdf(file_path):
    """Extract text from a .pdf file."""
    try:
        reader = PdfReader(file_path)
        text = []
        metadata = {}
        
        # Extract metadata if available
        pdf_info = reader.metadata
        if pdf_info:
            for key in pdf_info:
                if pdf_info[key] and str(key).lower() not in ['', 'none']:
                    metadata[str(key).lower()] = str(pdf_info[key])
        
        # Extract text from pages
        for page in reader.pages:
            page_text = page.extract_text()
            if page_text.strip():
                text.append(page_text)
        
        return {"text": "\n\n".join(text), "metadata": metadata}
    except Exception as e:
        print(f"Error extracting text from {file_path}: {e}")
        return {"text": f"ERROR: Could not extract text from {file_path}", "metadata": {}}

def extract_text_from_txt(file_path):
    """Extract text from a .txt file."""
    try:
        with open(file_path, 'r', encoding='utf-8', errors='replace') as file:
            text = file.read()
        return {"text": text, "metadata": {}}
    except Exception as e:
        print(f"Error extracting text from {file_path}: {e}")
        return {"text": f"ERROR: Could not extract text from {file_path}", "metadata": {}}

def extract_text_from_excel(file_path):
    """Extract text from Excel files (.xlsx, .xls)."""
    try:
        # Read all sheets
        excel_file = pd.ExcelFile(file_path)
        sheets_data = []
        
        for sheet_name in excel_file.sheet_names:
            df = pd.read_excel(excel_file, sheet_name)
            sheets_data.append(f"--- Sheet: {sheet_name} ---")
            
            # Convert DataFrame to string representation
            sheets_data.append(df.to_string(index=False))
        
        return {"text": "\n\n".join(sheets_data), "metadata": {"sheets": excel_file.sheet_names}}
    except Exception as e:
        print(f"Error extracting text from {file_path}: {e}")
        return {"text": f"ERROR: Could not extract text from {file_path}", "metadata": {}}

def determine_category(file_path, content):
    """Determine the appropriate category for a file based on path and content."""
    file_name = file_path.name.lower()
    file_path_str = str(file_path).lower()
    content_lower = content.lower() if content else ""
    
    # Check if it's a character profile
    if any(re.search(pattern, file_name) for pattern in CHARACTER_PROFILE_PATTERNS) or \
       "character" in file_path_str and "profile" in file_path_str or \
       re.search(r'char.*profile', content_lower, re.DOTALL):
        return "character_profiles"
    
    # Check if it's an in-game document
    for subcategory, patterns in IN_GAME_DOCS_PATTERNS.items():
        if any(re.search(pattern, file_name) for pattern in patterns) or \
           any(re.search(pattern, content_lower, re.DOTALL) for pattern in patterns):
            return f"in_game_documents/{subcategory}"
    
    # Check if it's a narrative
    if any(re.search(pattern, file_name) for pattern in NARRATIVE_PATTERNS) or \
       "scenario" in file_path_str or \
       any(re.search(pattern, content_lower, re.DOTALL) for pattern in NARRATIVE_PATTERNS):
        # Determine if it's complete or unfinished (simplistic approach)
        if "unfinished" in file_name or "unfinished" in content_lower or \
           "incomplete" in file_name or "incomplete" in content_lower:
            return "narratives/unfinished_scenarios"
        else:
            return "narratives/completed_scenarios"
    
    # Check if it's worldbuilding
    if any(re.search(pattern, file_name) for pattern in WORLDBUILDING_PATTERNS) or \
       "worldbuilding" in file_path_str or \
       any(re.search(pattern, content_lower, re.DOTALL) for pattern in WORLDBUILDING_PATTERNS):
        return "worldbuilding"
    
    # Check if it's OOC content
    for subcategory, patterns in OOC_PATTERNS.items():
        if any(re.search(pattern, file_name) for pattern in patterns) or \
           "ooc" in file_path_str or \
           any(re.search(pattern, content_lower, re.DOTALL) for pattern in patterns):
            return f"ooc_content/{subcategory}"
    
    # Check if it's timeline
    if any(re.search(pattern, file_name) for pattern in TIMELINE_PATTERNS) or \
       "timeline" in file_path_str or \
       any(re.search(pattern, content_lower, re.DOTALL) for pattern in TIMELINE_PATTERNS):
        return "timeline"
    
    # Default for unclassified content
    return "unknown"

def process_file(file_path, relative_source_dir):
    """Process a single file based on its type."""
    file_extension = file_path.suffix.lower()
    content_data = {"text": "", "metadata": {}}
    
    try:
        # Extract text based on file type
        if file_extension in ['.docx', '.doc']:
            content_data = extract_text_from_docx(file_path)
        elif file_extension == '.pdf':
            content_data = extract_text_from_pdf(file_path)
        elif file_extension == '.txt' or file_extension == '.rtf':
            content_data = extract_text_from_txt(file_path)
        elif file_extension in ['.xlsx', '.xls']:
            content_data = extract_text_from_excel(file_path)
        else:
            # Skip files we can't process
            print(f"Skipping unsupported file type: {file_path}")
            file_counts["skipped"] += 1
            return
        
        # Add file metadata
        content_data["metadata"]["source_file"] = str(file_path)
        content_data["metadata"]["source_type"] = relative_source_dir
        content_data["metadata"]["extraction_date"] = datetime.now().isoformat()
        
        # Determine category
        category = determine_category(file_path, content_data["text"])
        
        # Update counts
        if category == "unknown":
            file_counts["unknown"] += 1
        elif "/" in category:
            main_cat, sub_cat = category.split("/")
            file_counts[main_cat][sub_cat] += 1
        else:
            file_counts[category] += 1
        
        # Create output file name
        output_filename = f"{file_path.stem}.txt"
        
        # Create output file path
        output_file = TARGET_DIR / category / output_filename
        output_file.parent.mkdir(parents=True, exist_ok=True)
        
        # Write text content and metadata to output file
        with open(output_file, "w", encoding="utf-8") as f:
            f.write(f"---\n")
            f.write(f"source: {content_data['metadata']['source_file']}\n")
            f.write(f"category: {category}\n")
            
            # Add other metadata
            for key, value in content_data["metadata"].items():
                if key not in ["source_file"]:
                    f.write(f"{key}: {value}\n")
            
            f.write(f"---\n\n")
            f.write(content_data["text"])
        
        print(f"Processed: {file_path} → {output_file}")
        processed_files[str(file_path)] = str(output_file)
        
    except Exception as e:
        print(f"Error processing {file_path}: {e}")
        print(traceback.format_exc())
        file_counts["error"] += 1

def process_directory(source_dir, relative_source_name):
    """Process all files in a directory recursively."""
    # Skip the External Stuff folder and any paths containing "External"
    if "External Stuff" in str(source_dir) or "External" in str(source_dir):
        print(f"Skipping External folder: {source_dir}")
        return
        
    for item in source_dir.iterdir():
        if item.is_file():
            # Skip non-document files and very large files
            if item.suffix.lower() in ['.docx', '.doc', '.pdf', '.txt', '.rtf', '.xlsx', '.xls']:
                if item.stat().st_size < 15 * 1024 * 1024:  # Skip files larger than 15MB
                    process_file(item, relative_source_name)
                else:
                    print(f"Skipping large file: {item}")
                    file_counts["skipped"] += 1
        elif item.is_dir() and not item.name.startswith('.'):
            process_directory(item, relative_source_name)

def generate_metadata_summary():
    """Generate metadata summary for the processed files."""
    metadata = {
        "processing_date": datetime.now().isoformat(),
        "file_counts": file_counts,
        "processed_files_count": len(processed_files),
        "categories": {
            "character_profiles": [f for f in processed_files.values() if "/character_profiles/" in f],
            "in_game_documents": {
                "laws_policies": [f for f in processed_files.values() if "/in_game_documents/laws_policies/" in f],
                "intelligence": [f for f in processed_files.values() if "/in_game_documents/intelligence/" in f],
                "technical": [f for f in processed_files.values() if "/in_game_documents/technical/" in f],
                "diplomatic": [f for f in processed_files.values() if "/in_game_documents/diplomatic/" in f]
            },
            "narratives": {
                "completed_scenarios": [f for f in processed_files.values() if "/narratives/completed_scenarios/" in f],
                "unfinished_scenarios": [f for f in processed_files.values() if "/narratives/unfinished_scenarios/" in f]
            },
            "worldbuilding": [f for f in processed_files.values() if "/worldbuilding/" in f],
            "ooc_content": {
                "rules": [f for f in processed_files.values() if "/ooc_content/rules/" in f],
                "admin": [f for f in processed_files.values() if "/ooc_content/admin/" in f],
                "templates": [f for f in processed_files.values() if "/ooc_content/templates/" in f]
            },
            "timeline": [f for f in processed_files.values() if "/timeline/" in f],
            "unknown": [f for f in processed_files.values() if "/unknown/" in f],
        }
    }
    
    # Write metadata to a JSON file
    metadata_file = TARGET_DIR / "metadata" / "processing_summary.json"
    metadata_file.parent.mkdir(parents=True, exist_ok=True)
    
    with open(metadata_file, "w", encoding="utf-8") as f:
        json.dump(metadata, f, indent=2)
    
    print(f"Metadata summary written to {metadata_file}")
    
    # Generate a Markdown report for human readability
    md_file = TARGET_DIR / "metadata" / "processing_report.md"
    
    with open(md_file, "w", encoding="utf-8") as f:
        f.write("# EOTIR LLM Data Processing Report\n\n")
        f.write(f"**Processing Date:** {metadata['processing_date']}\n\n")
        
        f.write("## File Processing Summary\n\n")
        f.write(f"Total processed files: {metadata['processed_files_count']}\n\n")
        
        f.write("| Category | Count |\n")
        f.write("|----------|-------|\n")
        f.write(f"| Character Profiles | {file_counts['character_profiles']} |\n")
        
        for subcat, count in file_counts['in_game_documents'].items():
            f.write(f"| In-Game Documents: {subcat.replace('_', ' ').title()} | {count} |\n")
        
        for subcat, count in file_counts['narratives'].items():
            f.write(f"| Narratives: {subcat.replace('_', ' ').title()} | {count} |\n")
        
        f.write(f"| Worldbuilding | {file_counts['worldbuilding']} |\n")
        
        for subcat, count in file_counts['ooc_content'].items():
            f.write(f"| OOC Content: {subcat.replace('_', ' ').title()} | {count} |\n")
        
        f.write(f"| Timeline | {file_counts['timeline']} |\n")
        f.write(f"| Unknown | {file_counts['unknown']} |\n")
        f.write(f"| Skipped | {file_counts['skipped']} |\n")
        f.write(f"| Error | {file_counts['error']} |\n\n")
        
        f.write("## Next Steps for LLM Training\n\n")
        f.write("1. **Review categorized files** - Check the categorization and make any manual adjustments\n")
        f.write("2. **Data cleaning** - Review extracted text for any formatting issues\n")
        f.write("3. **Dataset preparation** - Convert the text files to your LLM training format\n")
        f.write("4. **Training and validation split** - Divide the data into training and validation sets\n")
        f.write("5. **Train your model** - Proceed with training your LLM using this prepared dataset\n")
    
    print(f"Human-readable report written to {md_file}")

def create_training_format_examples():
    """Create examples of training data format for different use cases."""
    examples_dir = TARGET_DIR / "metadata" / "training_examples"
    examples_dir.mkdir(parents=True, exist_ok=True)
    
    # Example 1: Character profile generation
    with open(examples_dir / "character_profile_example.jsonl", "w", encoding="utf-8") as f:
        f.write('{"prompt": "Generate a character profile for a naval officer in the Imperial Republic navy", "completion": "# CHARACTER PROFILE\\n\\n**Name**: Commander Elara Voss\\n\\n**Rank**: Commander, Imperial Republic Navy\\n\\n**Age**: 34\\n\\n**Homeworld**: Corellia\\n\\n**Physical Description**:\\nHeight: 5\'8\\"\\nBuild: Athletic, toned\\nHair: Black, kept in a regulation bun while on duty\\nEyes: Steel grey\\nDistinguishing Features: Small scar above right eyebrow from combat training\\n\\n**Personality**:\\nDisciplined and meticulous, Commander Voss approaches her duties with unwavering dedication. She maintains a professional demeanor at all times, though those who serve under her know she values loyalty and protects her crew fiercely. Voss is calculating in battle situations, preferring strategy over brute force. She has a dry, subtle sense of humor that emerges rarely in appropriate company.\\n\\n**Background**:\\nBorn to a family with three generations of naval service, Elara Voss was raised on Corellian fleet stories. She graduated top of her class at the Imperial Republic Naval Academy, specializing in tactical operations. Her first posting was aboard the cruiser *Vigilant* as a junior tactical officer, where she distinguished herself during the Outer Rim pacification campaigns. After a series of promotions for exemplary service, she now commands the destroyer *Adamant*.\\n\\n**Skills & Specialties**:\\n- Advanced tactical planning and battlefield assessment\\n- Ship-to-ship combat expertise\\n- Personnel management and leadership\\n- Fluent in Basic, Corellian, and Durese languages\\n- Hand-to-hand combat training\\n\\n**Service Record**:\\n- 3265 IE: Graduated Imperial Republic Naval Academy\\n- 3265-3268 IE: Junior Tactical Officer, IR Cruiser *Vigilant*\\n- 3268-3271 IE: Senior Tactical Officer, IR Battleship *Implacable*\\n- 3271-3274 IE: Executive Officer, IR Destroyer *Formidable*\\n- 3274-Present: Commander, IR Destroyer *Adamant*\\n\\n**Current Assignment**:\\nCommanding officer of the IR Destroyer *Adamant*, assigned to the 5th Fleet, Core Defense Division.\\n\\n**Notable Achievements**:\\n- Imperial Commendation for Valor during the Battle of Ralltiir\\n- Strategic Operations Excellence Medal\\n- Three-time recipient of the Fleet Commander\'s Citation\\n\\n**Relationships**:\\n- Mentor: Admiral Kastor Vane (Ret.)\\n- Rival: Commander Thaddeus Crane, IR Destroyer *Resolute*\\n- Friend: Captain Mira Drayson, IR Intelligence Division\\n\\n**Personal Motivations**:\\nCommander Voss seeks to uphold order and stability in the galaxy through disciplined naval service. She believes in the ideals of the Imperial Republic and sees herself as a protector of its citizens. Privately, she aspires to fleet command, motivated by her family\'s legacy and her own ambition for excellence."}\n')
    
    # Example 2: Document generation
    with open(examples_dir / "document_generation_example.jsonl", "w", encoding="utf-8") as f:
        f.write('{"prompt": "Generate an Imperial Republic Senate resolution on establishing trade relations with the Outer Rim territories", "completion": "# IMPERIAL REPUBLIC SENATE RESOLUTION 3279-C\\n\\n## ESTABLISHMENT OF PREFERENTIAL TRADE RELATIONS WITH THE OUTER RIM TERRITORIES\\n\\n*Proposed by the Committee on Interstellar Commerce and Economic Development*\\n\\n**WHEREAS**, the Imperial Republic seeks to promote economic prosperity and stability throughout all regions under its governance;\\n\\n**WHEREAS**, the Outer Rim Territories represent a significant untapped market and resource base that could enhance the overall economic strength of the Imperial Republic;\\n\\n**WHEREAS**, increased economic integration between the Core Worlds and Outer Rim Territories will foster political stability and reduce separatist sentiments;\\n\\n**WHEREAS**, the development of robust trade routes to the Outer Rim will create new opportunities for Imperial Republic citizens and businesses;\\n\\n**NOW, THEREFORE, BE IT RESOLVED** by the Imperial Senate of the Imperial Republic:\\n\\n**SECTION 1. ESTABLISHMENT OF THE OUTER RIM TRADE INITIATIVE**\\n\\nThe Senate hereby establishes the Outer Rim Trade Initiative (ORTI), which shall be administered by the Ministry of Commerce and Industry in coordination with the Ministry of Defense to ensure secure trade routes.\\n\\n**SECTION 2. TRADE INCENTIVES AND TARIFF REDUCTIONS**\\n\\n(a) All goods and services originating from Imperial Republic-aligned worlds in the Outer Rim Territories shall receive a 30% reduction in standard import tariffs for a period of five (5) standard years.\\n\\n(b) Imperial Republic corporations establishing manufacturing or resource extraction operations on Outer Rim worlds shall receive tax incentives as specified in Schedule A attached to this resolution.\\n\\n(c) The Ministry of Commerce shall establish expedited customs procedures for vessels registered to Imperial Republic citizens operating on designated Outer Rim trade routes.\\n\\n**SECTION 3. INFRASTRUCTURE DEVELOPMENT**\\n\\n(a) The sum of 50 billion credits is hereby appropriated for the expansion and reinforcement of hyperspace lanes connecting the Core Worlds to the Outer Rim Territories.\\n\\n(b) The Imperial Republic Navy is authorized to establish additional patrol stations along these trade routes to ensure the safety of commercial vessels.\\n\\n(c) The Ministry of Infrastructure shall prioritize the development of spaceport facilities on key Outer Rim worlds as designated in Schedule B attached to this resolution.\\n\\n**SECTION 4. OVERSIGHT AND REPORTING**\\n\\n(a) The Ministry of Commerce shall submit quarterly reports to the Senate Committee on Interstellar Commerce regarding the implementation and effects of this resolution.\\n\\n(b) The Imperial Economic Advisory Board shall conduct an annual review of the economic impact of these measures and make recommendations for adjustments as necessary.\\n\\n**SECTION 5. EFFECTIVE DATE**\\n\\nThis resolution shall take effect thirty (30) standard days after its passage by the Senate and approval by the Supreme Ruler.\\n\\n---\\n\\nAdopted by the Imperial Senate on the 15th day of the 4th month, 3279 Imperial Era.\\n\\n---\\n\\n*[Official Seal of the Imperial Senate]*\\n\\n---\\n\\nSenate President Marcellus Thorne\\n\\nChair, Committee on Interstellar Commerce\\nSenator Vora Kesh\\n\\nRecorded by:\\nSenate Secretary Livia Octavian"}\n')
    
    # Example 3: Scenario continuation
    with open(examples_dir / "scenario_continuation_example.jsonl", "w", encoding="utf-8") as f:
        scenario_prompt = '"Continue this unfinished scenario: \\n\\n'
        scenario_prompt += 'Admiral Kendra Voss stood on the bridge of the Imperial flagship Indomitable, staring at the tactical display. '
        scenario_prompt += 'The rebel fleet had emerged from hyperspace at the edge of the Coreward system, their formation suggesting they had '
        scenario_prompt += 'received intelligence about the Empire\'s defensive positions. \\\"How did they know?\\\" she whispered. '
        scenario_prompt += 'Captain Merrick approached, his expression grim. \\\"Admiral, there\'s something you need to see,\\\" '
        scenario_prompt += 'he said, handing her a datapad containing"'
        
        scenario_completion = '" classified communications intercepts.\\n\\n'
        scenario_completion += 'Admiral Voss\'s eyes narrowed as she scanned the contents. The intercepts revealed a series of encrypted transmissions '
        scenario_completion += 'between an Imperial communications relay in the Coreward sector and a known rebel outpost in the Mid Rim. '
        scenario_completion += 'The pattern was unmistakable - someone with high-level clearance was feeding information to the rebellion.\\n\\n'
        scenario_completion += '\\\"How long has this been happening?\\\" she demanded, her voice dangerously quiet.\\n\\n'
        scenario_completion += '\\\"Based on our analysis, approximately three months, Admiral,\\\" Captain Merrick replied. '
        scenario_completion += '\\\"Imperial Intelligence has been tracking the anomalies, but this is the first concrete evidence of a serious '
        scenario_completion += 'security breach. The traitor must be someone with access to our defense grid planning.\\\"\\n\\n'
        scenario_completion += 'Voss looked back at the tactical display. The rebel fleet - fifteen capital ships and their escorts - had formed '
        scenario_completion += 'into an attack position that perfectly exploited the gap in the Imperial defense network created by the recent '
        scenario_completion += 'redeployment of the 5th Fleet to quell uprisings in the Outer Rim.\\n\\n'
        scenario_completion += '\\\"They\'re not just aware of our positions,\\\" she realized. \\\"They know we\'re vulnerable right now.\\\"\\n\\n'
        scenario_completion += 'The bridge comms officer called out, \\\"Admiral, we\'re receiving a priority transmission from Imperial Center!\\\"\\n\\n'
        scenario_completion += '\\\"On screen,\\\" Voss ordered.\\n\\n'
        scenario_completion += 'The stern face of Grand Moff Tiberius appeared. \\\"Admiral Voss, we have reason to believe your Executive Officer, '
        scenario_completion += 'Commander Drayson, has been compromised. Imperial Intelligence has traced the security breach to command codes '
        scenario_completion += 'assigned to her station. She is to be detained immediately for questioning.\\\"\\n\\n'
        scenario_completion += 'Voss felt the blood drain from her face. Elara Drayson had been her trusted XO for over three years, a decorated '
        scenario_completion += 'officer with an impeccable service record. More importantly, she wasn\'t even on the bridge - Drayson had requested '
        scenario_completion += 'permission to inspect the starboard weapons batteries thirty minutes ago.\\n\\n'
        scenario_completion += '\\\"Security teams to starboard weapons control,\\\" Voss ordered immediately. '
        scenario_completion += '\\\"And lock down all escape pods and shuttlecraft.\\\"\\n\\n'
        scenario_completion += 'As the security alert klaxons began to sound throughout the massive Star Destroyer, the tactical officer shouted, '
        scenario_completion += '\\\"Admiral! Rebel fleet is accelerating to attack speed!\\\"\\n\\n'
        scenario_completion += 'Voss made her decision instantly. \\\"Battle stations! All power to forward shields. Inform all captains to execute '
        scenario_completion += 'Defense Pattern Epsilon immediately.\\\" She turned to Captain Merrick. \\\"Find Drayson. I want her alive.\\\"\\n\\n'
        scenario_completion += 'As the bridge erupted into controlled chaos, Voss stared at the approaching rebel fleet. The traitor had given them '
        scenario_completion += 'an opportunity, but the Imperial Republic Navy had prevailed against worse odds. Her mind raced through contingencies, '
        scenario_completion += 'evaluating and discarding tactical options.\\n\\n'
        scenario_completion += '\\\"Communications, transmit to all ships: The rebellion believes they have the advantage of surprise. Let\'s show them '
        scenario_completion += 'why the Imperial Republic has endured for a thousand years. Execute flanking maneuver through the Coreward asteroid field - '
        scenario_completion += 'the heavy cruisers won\'t be able to follow. We\'ll divide their force and eliminate them piecemeal.\\\"\\n\\n'
        scenario_completion += 'As orders were relayed and the massive Imperial fleet began its intricate dance through space, Voss received word that '
        scenario_completion += 'Commander Drayson had been cornered in Shuttle Bay 4, attempting to launch a diplomatic courier vessel.\\n\\n'
        scenario_completion += '\\\"I\'ll deal with betrayal later,\\\" Voss thought grimly as the first rebel cruiser came into weapons range. '
        scenario_completion += '\\\"For now, we have a battle to win.\\\"\\n\\n'
        scenario_completion += 'The forward batteries of the Indomitable opened fire, brilliant green turbolaser bolts streaking across the void '
        scenario_completion += 'toward the rebel vanguard. The battle for the Coreward system had begun."'
        
        f.write('{"prompt": ' + scenario_prompt + ', "completion": ' + scenario_completion + '}\n')
    
    # Example 4: Timeline integration
    with open(examples_dir / "timeline_integration_example.jsonl", "w", encoding="utf-8") as f:
        f.write('{"prompt": "Create a timeline entry for the following event:\\n\\nOn the 23rd day of the 9th month, 3281 Imperial Era, Supreme Commander Arlen Vexx led Imperial forces in the successful defense of the Coreward shipyards against a major rebel attack. Despite being outnumbered, Vexx\'s innovative tactical maneuvers resulted in the destruction of three rebel capital ships and the capture of rebel commander Sylas Krane.", "completion": "# IMPERIAL TIMELINE ENTRY\\n\\n## 3281.9.23 IE: THE BATTLE OF COREWARD SHIPYARDS\\n\\n**Classification**: Major Military Engagement\\n**Location**: Coreward System Shipyards, Core Sector\\n**Outcome**: Decisive Imperial Victory\\n\\n### SUMMARY:\\nImperial forces under Supreme Commander Arlen Vexx successfully defended the strategically vital Coreward Shipyards against a large-scale rebel assault. The rebel fleet, consisting of two heavy cruisers, four frigates, and multiple fighter squadrons, launched a surprise attack with the apparent goal of destroying the Imperial Republic\'s newest Sovereign-class Star Destroyer under construction at the facility.\\n\\nDespite being outnumbered, Supreme Commander Vexx employed innovative tactical maneuvers that lured the rebel forces into a prepared defensive grid. The battle resulted in the destruction of three rebel capital ships and minimal damage to the shipyards. Most significantly, rebel commander Sylas Krane was captured during the engagement when his command vessel was disabled by concentrated ion cannon fire.\\n\\n### STRATEGIC SIGNIFICANCE:\\n- Preserved the Coreward Shipyards\' production capacity for the Imperial Republic Fleet\\n- Captured high-ranking rebel commander Sylas Krane, providing valuable intelligence on rebel operations\\n- Demonstrated the tactical superiority of Imperial forces even when numerically disadvantaged\\n- Secured continued construction of the Sovereign-class Star Destroyer program\\n\\n### IMPERIAL REPUBLIC FORCES:\\n- Task Force Adamant (3 Star Destroyers, 6 light cruisers, 12 corvettes)\\n- Shipyard Defense Grid (orbital weapon platforms, tractor beam arrays)\\n- 2 TIE fighter wings\\n\\n### CASUALTIES:\\n- Imperial: 1 corvette lost, moderate damage to 2 light cruisers, 27 TIE fighters destroyed\\n- Rebel: 3 capital ships destroyed, 2 frigates severely damaged, estimated 1,200 casualties\\n\\n### COMMENDATIONS:\\n- Supreme Commander Arlen Vexx awarded the Star of the Imperial Republic for tactical brilliance\\n- Captain Elara Thorne received the Medal of Valor for defense of the primary construction dock\\n\\n### FOLLOW-UP ACTIONS:\\n- Interrogation of Commander Krane by Imperial Intelligence\\n- Security protocols at all Imperial shipyards reviewed and enhanced\\n- Accelerated completion schedule for the Sovereign-class Star Destroyer program"}\n')

def main():
    """Main execution function."""
    print("Starting EOTIR LLM training data extraction...")
    
    # Process the Novels directory
    print(f"\nProcessing {SOURCE_NOVELS_DIR}...")
    process_directory(SOURCE_NOVELS_DIR, "EOTIR Novels")
    
    # Process the RPG directory
    print(f"\nProcessing {SOURCE_RPG_DIR}...")
    process_directory(SOURCE_RPG_DIR, "EOTIR RPG")
    
    # Generate metadata summary
    print("\nGenerating metadata summary...")
    generate_metadata_summary()
    
    # Create training format examples
    print("\nCreating training format examples...")
    create_training_format_examples()
    
    print("\nData processing complete! Check the LLM directory for the processed files.")
    print(f"Total files processed: {sum(file_counts[cat] if isinstance(file_counts[cat], int) else sum(file_counts[cat].values()) for cat in file_counts)}")
    print(f"Files with errors: {file_counts['error']}")
    print(f"Files skipped: {file_counts['skipped']}")

if __name__ == "__main__":
    main()
